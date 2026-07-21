"""DOCX-Handler (python-docx).

Iteriert über Body-Paragraphen, Tabellen-Zellen und Header/Footer. Pro
Paragraph wird der gesamte Text auf einmal transformiert; Original-Runs
werden zu *einem* Run pro Paragraph zusammengefasst, um exakte
Pseudonym-Ersetzung zu garantieren — Run-Formatierung des ersten Runs
bleibt erhalten, weitere Runs werden gelöscht. Komplexe Inline-Formate
(z. B. fett-mitten-im-Wort) gehen damit ggf. verloren — ein bewusster
Trade-off zugunsten korrekter Anonymisierung.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

from pseudokrat.formats.base import (
    FormatProcessResult,
    TextTransform,
    derive_default_output,
    validate_office_archive,
)

if TYPE_CHECKING:  # pragma: no cover
    from docx.text.paragraph import Paragraph


class DocxHandler:
    name = "docx"
    suffixes: tuple[str, ...] = (".docx",)

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.suffixes

    def default_output_path(self, input_path: Path, suffix: str = "anon") -> Path:
        return derive_default_output(input_path, suffix=suffix)

    @staticmethod
    def _transform_paragraph(paragraph: Paragraph, transform: TextTransform) -> bool:
        """Liefert ``True``, wenn der Paragraph eine nicht-leere Transformation hatte."""
        original = paragraph.text
        if not original:
            return False
        new_text = transform(original)
        if new_text == original:
            return False
        # Schreibe Ergebnis in den ersten Run; entferne alle weiteren.
        runs = paragraph.runs
        if not runs:
            paragraph.add_run(new_text)
            return True
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
        return True

    def process(
        self,
        input_path: Path,
        output_path: Path,
        transform: TextTransform,
    ) -> FormatProcessResult:
        from docx import Document  # lokaler Import — schwere Lib
        from docx.oxml.ns import qn

        validate_office_archive(input_path)
        document = Document(str(input_path))
        processed = 0
        skipped = 0

        def _handle(paragraph: Paragraph) -> None:
            nonlocal processed, skipped
            if paragraph.text:
                if self._transform_paragraph(paragraph, transform):
                    processed += 1
                else:
                    skipped += 1

        def _handle_tables(tables: object) -> None:
            for table in tables:  # type: ignore[attr-defined]
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _handle(para)
                        _handle_tables(cell.tables)  # verschachtelte Tabellen

        # 1) Strukturierter Pass (Run-übergreifende Namen korrekt ersetzen):
        #    Body, (verschachtelte) Tabellen, ALLE Kopf-/Fusszeilen-Varianten.
        for para in document.paragraphs:
            _handle(para)
        _handle_tables(document.tables)
        for section in document.sections:
            for container in (
                section.header,
                section.footer,
                section.first_page_header,
                section.first_page_footer,
                section.even_page_header,
                section.even_page_footer,
            ):
                for para in container.paragraphs:
                    _handle(para)
                _handle_tables(container.tables)

        # 2) XML-Sweep über versteckte Kanäle: Kommentare, Fuss-/Endnoten,
        #    Textfelder (w:txbxContent) und Tracked-Changes-Löschungen
        #    (w:delText). Bereits anonymisierter Text ist PII-frei -> erneutes
        #    Anwenden ist ein No-op (idempotent).
        text_tags = {qn("w:t"), qn("w:delText")}
        roots: list[object] = [document.element]
        for rel in document.part.rels.values():
            try:
                target = rel.target_part
            except (ValueError, AttributeError):
                continue  # externe Relationships (Hyperlinks) o. ä.
            element = getattr(target, "element", None)
            if element is not None and hasattr(element, "iter"):
                roots.append(element)
        for root in roots:
            for el in root.iter():  # type: ignore[attr-defined]
                if el.tag in text_tags and el.text:
                    new_text = transform(el.text)
                    if new_text != el.text:
                        el.text = new_text
                        processed += 1

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))

        return FormatProcessResult(
            input_path=input_path,
            output_path=output_path,
            segments_processed=processed,
            segments_skipped=skipped,
        )

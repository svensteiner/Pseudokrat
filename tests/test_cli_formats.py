"""End-to-end CLI-Tests für strukturierte Dateiformate."""

from __future__ import annotations

from pathlib import Path

import pytest

from pseudokrat.cli import main


@pytest.fixture(autouse=True)
def _env(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    monkeypatch.setenv("PSEUDOKRAT_DATA_DIR", str(tmp_path / "data"))
    monkeypatch.setenv("PSEUDOKRAT_DISABLE_ML", "1")
    monkeypatch.setenv("PSEUDOKRAT_PASSWORD", "pw")
    return tmp_path


def test_cli_csv_roundtrip(_env: Path) -> None:
    inp = _env / "salden.csv"
    out = _env / "salden.anon.csv"
    back = _env / "salden.deanon.csv"
    inp.write_text(
        "Mandant;Saldo\nHofer Bau GmbH;1000\nHofer Bau GmbH;2000\n",
        encoding="utf-8",
    )

    rc = main(["anonymize", "--profile", "csv1", "-i", str(inp), "-o", str(out), "--no-ml"])
    assert rc == 0
    anonymized = out.read_text(encoding="utf-8")
    assert "<COMPANY_001>" in anonymized
    assert "Hofer Bau GmbH" not in anonymized

    rc2 = main(["deanonymize", "--profile", "csv1", "-i", str(out), "-o", str(back)])
    assert rc2 == 0
    restored = back.read_text(encoding="utf-8")
    assert "Hofer Bau GmbH" in restored


def test_cli_docx_roundtrip(_env: Path) -> None:
    from docx import Document

    inp = _env / "vertrag.docx"
    out = _env / "vertrag.anon.docx"

    doc = Document()
    doc.add_paragraph("Vertrag mit Hofer Bau GmbH über IBAN AT611904300234573201.")
    doc.save(str(inp))

    rc = main(["anonymize", "--profile", "docx1", "-i", str(inp), "-o", str(out), "--no-ml"])
    assert rc == 0

    anonymized = Document(str(out))
    full = " ".join(p.text for p in anonymized.paragraphs)
    assert "Hofer Bau GmbH" not in full
    assert "<COMPANY_001>" in full
    assert "<IBAN_001>" in full


def test_cli_xlsx_roundtrip(_env: Path) -> None:
    from openpyxl import Workbook, load_workbook

    inp = _env / "salden.xlsx"
    out = _env / "salden.anon.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["Mandant", "Saldo"])
    ws.append(["Hofer Bau GmbH", 1000])
    ws.append(["Hofer Bau GmbH", 2000])
    wb.save(str(inp))

    rc = main(["anonymize", "--profile", "xlsx1", "-i", str(inp), "-o", str(out), "--no-ml"])
    assert rc == 0

    wb_out = load_workbook(str(out))
    ws_out = wb_out.active
    assert ws_out["A2"].value == ws_out["A3"].value
    assert ws_out["A2"].value.startswith("<COMPANY_")
    assert ws_out["B2"].value == 1000  # Zahl unverändert


def test_cli_default_output_path_when_omitted(_env: Path) -> None:
    inp = _env / "memo.csv"
    inp.write_text("Mandant\nHofer Bau GmbH\n", encoding="utf-8")
    rc = main(["anonymize", "--profile", "p", "-i", str(inp), "--no-ml"])
    assert rc == 0
    assert (_env / "memo.anon.csv").exists()


def test_cli_unsupported_extension_falls_back_to_text_path(_env: Path) -> None:
    """Eine Endung ohne registrierten Handler (z. B. ``.zip``) geht über
    den Text-Pfad. UTF-8-Decoding-Fehler quittiert die CLI mit Exception
    bzw. Exit-Code; hier prüfen wir nur, dass ``_has_handler`` für
    unbekannte Endungen ``False`` liefert."""
    from pseudokrat.cli import _has_handler

    assert _has_handler(_env / "anything.zip") is False
    assert _has_handler(_env / "no_ext") is False

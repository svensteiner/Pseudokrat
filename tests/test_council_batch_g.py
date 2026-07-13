"""Tests fuer Council-Batch-G (#8): DOCX versteckte Kanaele —
Tracked-Changes-Loeschungen (w:delText) und verschachtelte Tabellen."""

from __future__ import annotations

from pathlib import Path

from pseudokrat.formats.docx_handler import DocxHandler


def _mask(t: str) -> str:
    return t.replace("Hankook", "<COMPANY_001>")


def test_tracked_deletion_anonymized(tmp_path: Path) -> None:
    from docx import Document
    from docx.oxml import OxmlElement

    doc = Document()
    para = doc.add_paragraph("Sichtbarer Text.")
    run = OxmlElement("w:r")
    del_text = OxmlElement("w:delText")
    del_text.text = "Hankook Reifen"
    run.append(del_text)
    para._p.append(run)
    src = tmp_path / "in.docx"
    doc.save(str(src))

    out = tmp_path / "out.docx"
    DocxHandler().process(src, out, transform=_mask)

    xml = Document(str(out)).element.xml
    assert "Hankook" not in xml  # geloeschter Text (Tracked Changes) ist weg
    assert "COMPANY_001" in xml


def test_nested_table_anonymized(tmp_path: Path) -> None:
    from docx import Document

    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    inner = outer.rows[0].cells[0].add_table(rows=1, cols=1)
    inner.rows[0].cells[0].text = "Hankook GmbH"
    src = tmp_path / "nested.docx"
    doc.save(str(src))

    out = tmp_path / "nested.out.docx"
    DocxHandler().process(src, out, transform=_mask)

    xml = Document(str(out)).element.xml
    assert "Hankook" not in xml


def test_body_still_anonymized(tmp_path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Bericht der Hankook GmbH ueber das Jahr.")
    src = tmp_path / "b.docx"
    doc.save(str(src))
    out = tmp_path / "b.out.docx"
    DocxHandler().process(src, out, transform=_mask)
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Hankook" not in text
    assert "<COMPANY_001>" in text

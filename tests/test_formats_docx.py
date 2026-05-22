"""Tests für DOCX-Handler — Phase-4-Vorzieher."""

from __future__ import annotations

from pathlib import Path

from pseudokrat.formats.docx_handler import DocxHandler


def _make_docx(
    path: Path, paragraphs: list[str], table_cells: list[list[str]] | None = None
) -> None:
    from docx import Document

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_cells:
        rows = len(table_cells)
        cols = len(table_cells[0])
        table = doc.add_table(rows=rows, cols=cols)
        for i, row in enumerate(table_cells):
            for j, val in enumerate(row):
                table.cell(i, j).text = val
    doc.save(str(path))


def test_docx_paragraph_anonymization(tmp_path: Path) -> None:
    inp = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_docx(
        inp,
        paragraphs=[
            "Schreiben an Hofer Bau GmbH.",
            "",  # leerer Paragraph
            "Konto IBAN AT611904300234573201.",
        ],
    )

    def transform(text: str) -> str:
        return text.replace("Hofer Bau GmbH", "<COMPANY_001>").replace(
            "AT611904300234573201", "<IBAN_001>"
        )

    result = DocxHandler().process(inp, out, transform=transform)
    assert result.segments_processed >= 2

    from docx import Document

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert any("<COMPANY_001>" in t for t in texts)
    assert any("<IBAN_001>" in t for t in texts)
    assert not any("Hofer Bau GmbH" in t for t in texts)


def test_docx_table_cells_processed(tmp_path: Path) -> None:
    inp = tmp_path / "table.docx"
    out = tmp_path / "table.anon.docx"
    _make_docx(
        inp,
        paragraphs=["Übersicht:"],
        table_cells=[["Mandant", "Saldo"], ["Hofer Bau GmbH", "1000"]],
    )

    def transform(text: str) -> str:
        return text.replace("Hofer Bau GmbH", "<X>")

    DocxHandler().process(inp, out, transform=transform)

    from docx import Document

    doc = Document(str(out))
    table_texts = [cell.text for row in doc.tables[0].rows for cell in row.cells]
    assert "<X>" in table_texts
    assert "Hofer Bau GmbH" not in table_texts


def test_docx_default_output_and_supports() -> None:
    h = DocxHandler()
    assert h.supports(Path("x.docx"))
    assert not h.supports(Path("x.txt"))
    assert h.default_output_path(Path("/tmp/x.docx")) == Path("/tmp/x.anon.docx")


def test_docx_unchanged_paragraph_is_skipped(tmp_path: Path) -> None:
    inp = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    _make_docx(inp, paragraphs=["Nichts Sensibles hier."])

    # Identity-Transform → segments_skipped = 1, segments_processed = 0
    result = DocxHandler().process(inp, out, transform=lambda t: t)
    assert result.segments_processed == 0
    assert result.segments_skipped == 1

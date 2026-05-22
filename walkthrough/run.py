"""End-to-End-Walkthrough für Pseudokrat — wie ein Steuerberater es täte.

Ablauf:
  01  Smoke: Modul + Version
  02  CLI --help spricht deutsch und listet alle Subcommands
  03  CLI anonymize via --text auf frischem Profil
  04  CLI deanonymize Round-Trip via Pipe
  05  CLI Datei-Pipeline TXT
  06  CLI Datei-Pipeline CSV
  07  CLI Datei-Pipeline DOCX
  08  CLI Datei-Pipeline XLSX (inkl. Formel-Konsistenz)
  09  Multi-Session-Konsistenz: 2x öffnen, gleiche Firma → gleicher Platzhalter
  10  Fuzzy-Merge: drei Schreibweisen → ein Pseudonym
  11  Audit-Log: verify + export, Hash-Chain valid
  12  Profil-Liste enthält alle drei angelegten Profile
  13  GUI headless: Live-Workflow + Datei-Workflow
  14  Falsches Passwort blockt Profil-Öffnen
  15  Unbekannter Platzhalter wird beim Deanonymisieren toleriert
  16  CLI Datei-Pipeline PDF (Round-Trip Anonymize→Deanonymize)
  17  Audit-Log PDF-Export (reportlab)
  18  Per-Profil-Mandanten-Pattern (§7): init mit --mandanten-pattern, Round-Trip

Exit-Code 0 = alle Schritte grün. Erster Fehler bricht ab.
"""

from __future__ import annotations

import io
import os
import shutil
import subprocess
import sys
import tempfile
import textwrap
import traceback
from contextlib import redirect_stdout, redirect_stderr
from pathlib import Path

# Frische Umgebung: kompletter Walkthrough nutzt ein einziges tmp-Verzeichnis
ROOT = Path(tempfile.mkdtemp(prefix="pseudokrat_walkthrough_"))
DATA = ROOT / "data"
FILES = ROOT / "files"
DATA.mkdir(parents=True, exist_ok=True)
FILES.mkdir(parents=True, exist_ok=True)

os.environ["PSEUDOKRAT_DATA_DIR"] = str(DATA)
os.environ["PSEUDOKRAT_DISABLE_ML"] = "1"
os.environ["PSEUDOKRAT_PASSWORD"] = "demo-pw-2026"
os.environ["QT_QPA_PLATFORM"] = "offscreen"

# Imports NACH env-Setup, damit Settings das DATA-Verzeichnis sehen.
from pseudokrat.cli import main as cli_main  # noqa: E402


def header(num: str, title: str) -> None:
    print(f"\n=== [{num}] {title} " + "=" * max(0, 60 - len(title)))


def run_cli(args: list[str], stdin: str | None = None) -> tuple[int, str, str]:
    """Führe die CLI im selben Prozess aus, mit aufgefangenem stdout/stderr."""
    out, err = io.StringIO(), io.StringIO()
    saved_stdin = sys.stdin
    if stdin is not None:
        sys.stdin = io.StringIO(stdin)
    rc = 1
    try:
        with redirect_stdout(out), redirect_stderr(err):
            try:
                rc = cli_main(args)
            except SystemExit as exc:
                rc = int(exc.code or 0)
    finally:
        sys.stdin = saved_stdin
    return rc, out.getvalue(), err.getvalue()


def assert_eq(actual: object, expected: object, label: str) -> None:
    if actual != expected:
        raise AssertionError(f"{label}: erwartet {expected!r}, war {actual!r}")


def assert_in(needle: str, haystack: str, label: str) -> None:
    if needle not in haystack:
        raise AssertionError(f"{label}: '{needle}' nicht in:\n{haystack[:500]}")


def assert_not_in(needle: str, haystack: str, label: str) -> None:
    if needle in haystack:
        raise AssertionError(f"{label}: '{needle}' war noch enthalten:\n{haystack[:500]}")


# --- Steps -----------------------------------------------------------------


def step_01_smoke() -> None:
    header("01", "Smoke: Modul + Version")
    import pseudokrat

    assert pseudokrat.__version__.startswith("0."), pseudokrat.__version__
    print(f"   pseudokrat == {pseudokrat.__version__}, Python {sys.version.split()[0]}")


def step_02_help() -> None:
    header("02", "CLI --help (deutsch, 4 Subcommands)")
    rc, out, _ = run_cli(["--help"])
    assert_eq(rc, 0, "Exit-Code")
    for sub in ("anonymize", "deanonymize", "profiles", "audit"):
        assert_in(sub, out, f"Subcommand '{sub}' fehlt im Help")
    assert_in("DACH", out, "DACH-Beschreibung fehlt")
    print("   /OK alle 4 Subcommands sichtbar")


def step_03_anonymize_text() -> None:
    header("03", "CLI anonymize via --text (neues Profil)")
    rc, out, err = run_cli(
        [
            "anonymize",
            "--profile",
            "Mandant Hofer",
            "--text",
            "Die Hofer Bau GmbH (UID ATU12345675) überweist 1.200 € auf AT611904300234573201.",
            "--no-ml",
        ]
    )
    assert_eq(rc, 0, "Exit-Code")
    assert_in("<COMPANY_001>", out, "COMPANY-Platzhalter fehlt")
    assert_in("<UID_001>", out, "UID-Platzhalter fehlt")
    assert_in("<IBAN_001>", out, "IBAN-Platzhalter fehlt")
    assert_in("1.200 €", out, "Betrag wurde fälschlich maskiert")
    print("   /OK Output:", out.strip().splitlines()[-1])


def step_04_deanonymize() -> None:
    header("04", "CLI deanonymize Round-Trip")
    anonymized = (
        "Die <COMPANY_001> sollte für <IBAN_001> einen Dauerauftrag einrichten."
    )
    rc, out, _ = run_cli(
        [
            "deanonymize",
            "--profile",
            "Mandant Hofer",
            "--text",
            anonymized,
        ]
    )
    assert_eq(rc, 0, "Exit-Code")
    assert_in("Hofer Bau GmbH", out, "Firma nicht rückübersetzt")
    assert_in("AT611904300234573201", out, "IBAN nicht rückübersetzt")
    print("   /OK Round-Trip Output:", out.strip().splitlines()[-1])


def step_05_txt_pipeline() -> None:
    header("05", "Datei-Pipeline TXT (-i memo.txt)")
    src = FILES / "memo.txt"
    out_path = FILES / "memo.anon.txt"
    src.write_text(
        "An Hofer Bau GmbH: Bitte beachten Sie die UID ATU12345675.\n",
        encoding="utf-8",
    )
    rc, _, err = run_cli(
        ["anonymize", "--profile", "Mandant Hofer", "-i", str(src), "--no-ml"]
    )
    assert_eq(rc, 0, "Exit-Code")
    assert out_path.exists(), "Standardausgabedatei fehlt"
    content = out_path.read_text(encoding="utf-8")
    assert_in("<COMPANY_001>", content, "TXT Anonymisat")
    assert_not_in("Hofer Bau GmbH", content, "TXT-Original-Sicht")
    print(f"   /OK schrieb {out_path.name}")


def step_06_csv_pipeline() -> None:
    header("06", "Datei-Pipeline CSV (Spaltenkonsistenz)")
    src = FILES / "salden.csv"
    out_path = FILES / "salden.anon.csv"
    src.write_text(
        "Mandant;Saldo\nHofer Bau GmbH;1000\nHofer Bau GmbH;2000\nMüller AG;5000\n",
        encoding="utf-8",
    )
    rc, _, _ = run_cli(
        [
            "anonymize",
            "--profile",
            "Mandant Hofer",
            "-i",
            str(src),
            "-o",
            str(out_path),
            "--no-ml",
        ]
    )
    assert_eq(rc, 0, "Exit-Code")
    rows = out_path.read_text(encoding="utf-8").splitlines()
    assert len(rows) == 4, f"Zeilenzahl falsch: {len(rows)}"
    assert rows[1].split(";")[0] == rows[2].split(";")[0], "Hofer-Konsistenz"
    assert rows[1].split(";")[0] != rows[3].split(";")[0], "Müller darf nicht gemerged sein"
    assert rows[1].split(";")[1] == "1000", "Saldo wurde maskiert"
    print("   /OK CSV-Konsistenz:", rows[1], "/", rows[3])


def step_07_docx_pipeline() -> None:
    header("07", "Datei-Pipeline DOCX")
    from docx import Document

    src = FILES / "vertrag.docx"
    out_path = FILES / "vertrag.anon.docx"
    doc = Document()
    doc.add_paragraph("Vertrag mit Hofer Bau GmbH über IBAN AT611904300234573201.")
    doc.add_paragraph("Frist: 30 Tage.")
    doc.save(str(src))

    rc, _, _ = run_cli(
        ["anonymize", "--profile", "Mandant Hofer", "-i", str(src), "--no-ml"]
    )
    assert_eq(rc, 0, "Exit-Code")
    if not out_path.exists():
        raise AssertionError(f"DOCX-Output fehlt: {out_path}")

    out_doc = Document(str(out_path))
    paragraphs = [p.text for p in out_doc.paragraphs]
    full = " | ".join(paragraphs)
    assert_in("<COMPANY_001>", full, "DOCX-Anonymisat")
    assert_in("<IBAN_001>", full, "DOCX-IBAN")
    assert_not_in("Hofer Bau GmbH", full, "DOCX-Original")
    print("   /OK DOCX-Paragraphen:", paragraphs)


def step_08_xlsx_pipeline() -> None:
    header("08", "Datei-Pipeline XLSX (Formel-Konsistenz)")
    from openpyxl import Workbook, load_workbook

    src = FILES / "salden.xlsx"
    out_path = FILES / "salden.anon.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Salden"
    ws.append(["Mandant", "Saldo"])
    ws.append(["Hofer Bau GmbH", 1000])
    ws.append(["Hofer Bau GmbH", 2000])
    ws.append(["Müller AG", 5000])
    ws["D2"] = '=SUMIF(A:A,"Hofer Bau GmbH",B:B)'
    wb.save(str(src))

    rc, _, _ = run_cli(
        ["anonymize", "--profile", "Mandant Hofer", "-i", str(src), "--no-ml"]
    )
    assert_eq(rc, 0, "Exit-Code")

    out_wb = load_workbook(str(out_path))
    out_ws = out_wb.active
    assert out_ws["A2"].value == out_ws["A3"].value, "Hofer-Konsistenz XLSX"
    assert out_ws["A2"].value != out_ws["A4"].value, "Müller darf nicht merge"
    assert out_ws["B2"].value == 1000, "Zahlen wurden geändert"
    formula = str(out_ws["D2"].value)
    assert formula.startswith("=SUMIF"), f"Formel-Struktur verloren: {formula}"
    assert "Hofer Bau GmbH" not in formula, "Original in Formel"
    assert "<COMPANY_" in formula, "Platzhalter in Formel"
    print(f"   /OK Formel danach: {formula}")


def step_09_multi_session() -> None:
    header("09", "Multi-Session-Konsistenz")
    # Schließe + öffne dasselbe Profil über zwei CLI-Calls; gleiche Firma → gleicher Tag
    rc1, out1, _ = run_cli(
        [
            "anonymize",
            "--profile",
            "Session-Test",
            "--text",
            "Erste Erwähnung: Hofer Bau GmbH.",
            "--no-ml",
        ]
    )
    first = [line for line in out1.splitlines() if "<COMPANY_" in line][0]
    rc2, out2, _ = run_cli(
        [
            "anonymize",
            "--profile",
            "Session-Test",
            "--text",
            "Zweite Erwähnung: Hofer Bau GmbH.",
            "--no-ml",
        ]
    )
    second = [line for line in out2.splitlines() if "<COMPANY_" in line][0]
    if "<COMPANY_001>" not in first or "<COMPANY_001>" not in second:
        raise AssertionError(f"Konsistenz fehlgeschlagen: {first!r} vs {second!r}")
    print(f"   /OK beide Calls bekamen <COMPANY_001>")


def step_10_fuzzy_merge() -> None:
    header("10", "Fuzzy-Merge dreier Schreibweisen")
    rc, out, _ = run_cli(
        [
            "anonymize",
            "--profile",
            "Fuzzy-Test",
            "--text",
            "Hofer Bau GmbH, Hofer-Bau GmbH und hofer bau GmbH sind dieselbe Firma.",
            "--no-ml",
        ]
    )
    assert_eq(rc, 0, "Exit-Code")
    line = next(line for line in out.splitlines() if "dieselbe Firma" in line)
    count = line.count("<COMPANY_001>")
    if count != 3:
        raise AssertionError(f"Erwartet 3x <COMPANY_001>, war {count} in:\n{line}")
    if "<COMPANY_002>" in line:
        raise AssertionError("Fuzzy-Merge versagte — <COMPANY_002> aufgetaucht")
    print(f"   /OK alle drei Schreibweisen → <COMPANY_001>")


def step_11_audit() -> None:
    header("11", "Audit verify + export")
    rc1, out1, _ = run_cli(["audit", "--profile", "Mandant Hofer", "verify"])
    assert_eq(rc1, 0, "Verify-Exit")
    assert_in("OK", out1, "Hash-Chain")

    rc2, out2, _ = run_cli(["audit", "--profile", "Mandant Hofer", "export"])
    assert_eq(rc2, 0, "Export-Exit")
    header_line = out2.splitlines()[0]
    assert_in("timestamp", header_line, "CSV-Header")
    print(f"   /OK Hash-Chain valid, CSV-Header: {header_line}")


def step_12_profiles_list() -> None:
    header("12", "profiles list zeigt alle Profile")
    rc, out, _ = run_cli(["profiles", "list"])
    assert_eq(rc, 0, "Exit-Code")
    for name in ("Mandant Hofer", "Session-Test", "Fuzzy-Test"):
        assert_in(name, out, f"Profil '{name}' fehlt")
    print(f"   /OK alle drei Profile gelistet")


def step_13_gui_headless() -> None:
    header("13", "GUI headless (offscreen)")
    from PySide6.QtWidgets import QApplication

    from pseudokrat.gui.main_window import MainWindow, build_application

    app = QApplication.instance() or build_application(["walkthrough"])

    win = MainWindow()
    try:
        # Live-Workflow
        win.profile_input.setText("GUI-Walk")
        win.password_input.setText("demo-pw-2026")
        win._open_profile()
        win.input_edit.setPlainText("Schreiben an Hofer Bau GmbH (IBAN AT611904300234573201).")
        win._anonymize()
        anonymized = win.output_edit.toPlainText()
        assert_in("<COMPANY_001>", anonymized, "GUI-COMPANY")
        assert_in("<IBAN_001>", anonymized, "GUI-IBAN")
        win.input_edit.setPlainText(anonymized)
        win._deanonymize()
        restored = win.output_edit.toPlainText()
        assert_in("Hofer Bau GmbH", restored, "GUI-Round-Trip")

        # Datei-Workflow
        src = FILES / "gui_brief.txt"
        src.write_text("Hofer Bau GmbH ist Mandant.", encoding="utf-8")
        win.file_list.add_path(src)
        win._anonymize_files()
        target = src.with_name("gui_brief.anon.txt")
        assert target.exists(), "Datei-Tab erzeugte kein Output"
        assert_in("<COMPANY_001>", target.read_text(encoding="utf-8"), "GUI-Datei-Pipeline")
        print("   /OK Live + Datei-Tab beide funktionsfähig")
    finally:
        win.close()


def step_14_wrong_password() -> None:
    header("14", "Falsches Passwort wird sauber abgewiesen")
    rc1, _, _ = run_cli(
        [
            "anonymize",
            "--profile",
            "Locked",
            "--password",
            "richtig",
            "--text",
            "x",
            "--no-ml",
        ]
    )
    assert_eq(rc1, 0, "Erst-Anlage")
    rc2, _, err2 = run_cli(
        [
            "anonymize",
            "--profile",
            "Locked",
            "--password",
            "falsch",
            "--text",
            "y",
            "--no-ml",
        ]
    )
    if rc2 == 0:
        raise AssertionError("Falsches Passwort hätte fehlschlagen müssen")
    assert_in("Passwort", err2 + " ", "Fehlermeldung")
    print(f"   /OK Exit-Code {rc2}, Fehler: {err2.strip().splitlines()[-1]}")


def step_15_unknown_placeholder() -> None:
    header("15", "Unbekannter Platzhalter beim Deanonymisieren")
    rc, out, err = run_cli(
        [
            "deanonymize",
            "--profile",
            "Mandant Hofer",
            "--text",
            "Da kommt <COMPANY_999> nicht vor.",
        ]
    )
    # rc=3 nach CLI-Konvention bei missing_placeholders
    if rc not in (0, 3):
        raise AssertionError(f"Unerwarteter Exit-Code {rc}")
    assert_in("<COMPANY_999>", out, "Unbekannter Platzhalter sollte unverändert bleiben")
    print(f"   /OK exit={rc}, Text: {out.strip().splitlines()[-1]}")


def step_16_pdf_pipeline() -> None:
    header("16", "Datei-Pipeline PDF (Round-Trip)")
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen.canvas import Canvas

    src = FILES / "brief.pdf"
    anon = FILES / "brief.anon.pdf"
    deanon = FILES / "brief.deanon.pdf"

    # Eingangs-PDF mit zwei Seiten erzeugen.
    c = Canvas(str(src), pagesize=A4)
    width, height = A4
    for idx, line in enumerate(
        [
            "Schreiben an Hofer Bau GmbH.",
            "IBAN AT611904300234573201, Frist: 30 Tage.",
        ]
    ):
        if idx > 0:
            c.showPage()
        c.setFont("Helvetica", 11)
        c.drawString(72, height - 72, line)
    c.save()

    rc, _, err = run_cli(
        ["anonymize", "--profile", "Mandant Hofer", "-i", str(src), "-o", str(anon)]
    )
    assert_eq(rc, 0, "PDF-Anonymize-Exit")
    assert anon.exists(), "Anonymisierte PDF fehlt"

    from pypdf import PdfReader

    anon_text = "\n".join((p.extract_text() or "") for p in PdfReader(str(anon)).pages)
    assert_in("<COMPANY_001>", anon_text, "PDF-COMPANY-Pseudonym")
    assert_in("<IBAN_001>", anon_text, "PDF-IBAN-Pseudonym")
    if "Hofer Bau GmbH" in anon_text:
        raise AssertionError("Original-Firma noch in anonymisierter PDF")

    rc2, _, _ = run_cli(
        ["deanonymize", "--profile", "Mandant Hofer", "-i", str(anon), "-o", str(deanon)]
    )
    assert_eq(rc2, 0, "PDF-Deanonymize-Exit")
    deanon_text = "\n".join((p.extract_text() or "") for p in PdfReader(str(deanon)).pages)
    assert_in("Hofer Bau GmbH", deanon_text, "PDF-Round-Trip-Firma")
    assert_in("AT611904300234573201", deanon_text, "PDF-Round-Trip-IBAN")
    print(f"   /OK PDF Round-Trip: <COMPANY_001>/<IBAN_001> ↔ Original wiederhergestellt")


def step_17_audit_pdf_export() -> None:
    header("17", "Audit-Log PDF-Export")
    out_pdf = FILES / "audit_hofer.pdf"
    rc, _, err = run_cli(
        [
            "audit",
            "--profile",
            "Mandant Hofer",
            "export",
            "--format",
            "pdf",
            "-o",
            str(out_pdf),
        ]
    )
    assert_eq(rc, 0, "Audit-PDF-Exit")
    assert out_pdf.exists(), "Audit-PDF fehlt"

    from pypdf import PdfReader

    text = "\n".join((p.extract_text() or "") for p in PdfReader(str(out_pdf)).pages)
    assert_in("Pseudokrat", text, "PDF-Title")
    assert_in("Mandant Hofer", text, "Profilname")
    assert_in("anonymize", text, "Audit-Eintrag")
    print(f"   /OK Audit-PDF generiert: {out_pdf.name}, {out_pdf.stat().st_size} Bytes")


def step_18_mandanten_pattern() -> None:
    header("18", "Per-Profil-Mandanten-Pattern (§7)")
    # init mit Pattern, dann Round-Trip
    rc1, out1, err1 = run_cli(
        [
            "init",
            "--profile",
            "Pattern-Test",
            "--mandanten-pattern",
            r"M-\d{5}",
        ]
    )
    assert_eq(rc1, 0, "init-Exit")
    assert_in("Mandanten-Pattern", out1, "init bestätigt Pattern")

    text = "Akte M-12345 betrifft Hofer Bau GmbH."
    rc2, out2, _ = run_cli(
        [
            "anonymize",
            "--profile",
            "Pattern-Test",
            "--text",
            text,
            "--no-ml",
        ]
    )
    assert_eq(rc2, 0, "Anonymize-Exit")
    assert_in("<MANDANT_NR_001>", out2, "Mandanten-Pattern wirkt")
    assert_not_in("M-12345", out2, "Mandantennummer noch sichtbar")

    # show-mandanten-pattern liest ohne Passwort
    rc3, out3, _ = run_cli(
        ["profiles", "show-mandanten-pattern", "--profile", "Pattern-Test"]
    )
    assert_eq(rc3, 0, "show-Exit")
    if r"M-\d{5}" not in out3:
        raise AssertionError(f"show-mandanten-pattern lieferte unerwartet: {out3!r}")

    # Round-Trip mit Deanonymize
    anon_line = next(line for line in out2.splitlines() if "<MANDANT_NR_001>" in line)
    rc4, out4, _ = run_cli(
        [
            "deanonymize",
            "--profile",
            "Pattern-Test",
            "--text",
            anon_line,
        ]
    )
    assert_eq(rc4, 0, "Deanon-Exit")
    assert_in("M-12345", out4, "Deanonymize stellt Mandantennummer wieder her")
    print(f"   /OK Pattern aktiv & Round-Trip: M-12345 ↔ <MANDANT_NR_001>")


# --- Driver ----------------------------------------------------------------

STEPS = [
    step_01_smoke,
    step_02_help,
    step_03_anonymize_text,
    step_04_deanonymize,
    step_05_txt_pipeline,
    step_06_csv_pipeline,
    step_07_docx_pipeline,
    step_08_xlsx_pipeline,
    step_09_multi_session,
    step_10_fuzzy_merge,
    step_11_audit,
    step_12_profiles_list,
    step_13_gui_headless,
    step_14_wrong_password,
    step_15_unknown_placeholder,
    step_16_pdf_pipeline,
    step_17_audit_pdf_export,
    step_18_mandanten_pattern,
]


def main() -> int:
    print(f"Walkthrough-Datenverzeichnis: {ROOT}")
    failed: list[tuple[str, str]] = []
    for step in STEPS:
        try:
            step()
        except Exception as exc:  # noqa: BLE001
            tb = "".join(traceback.format_exception(exc))
            print(f"\n[FAIL] {step.__name__}\n{tb}")
            failed.append((step.__name__, str(exc)))
            break  # erster Fehler -> abbrechen, damit man fokussiert nachbessern kann

    if failed:
        print(f"\n*** {len(failed)} Schritt(e) FEHLGESCHLAGEN ***")
        for name, msg in failed:
            print(f"   • {name}: {msg}")
        return 1

    print(f"\n*** ALLE {len(STEPS)} SCHRITTE GRÜN ***")
    # Aufräumen
    shutil.rmtree(ROOT, ignore_errors=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
